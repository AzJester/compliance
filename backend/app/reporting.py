from __future__ import annotations

import csv
import io
import re
from collections import Counter, defaultdict
from dataclasses import dataclass
from datetime import UTC, datetime
from typing import Final
from zoneinfo import ZoneInfo, ZoneInfoNotFoundError

from docx import Document as WordDocument
from docx.document import Document as WordDocumentType
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.table import Table, _Cell

DOCX_MEDIA_TYPE: Final = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
GAP_CSV_MEDIA_TYPE: Final = "text/csv; charset=utf-8"
REPORT_FINDING_LIMIT: Final = 25
REPORT_ACTION_LIMIT: Final = 50
_CONTENT_WIDTH_DXA: Final = 9360
_TABLE_INDENT_DXA: Final = 120
_XML_INVALID = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f\ud800-\udfff\ufffe\uffff]")
_FORMULA_PREFIXES: Final = ("=", "+", "-", "@")
_EXCEL_CELL_LIMIT: Final = 32_767
_STATUS_ORDER: Final = {"CONFLICT": 0, "MISSING": 1, "PARTIAL": 2}
_SECTION_ORDER: Final = {"M": 0, "L": 1, "C": 2, "H": 3}
_PRIORITY_LABEL: Final = {"CONFLICT": "P0", "MISSING": "P1", "PARTIAL": "P2"}
_STATUS_COLORS: Final = {
    "COVERED": "1F3A5F",
    "PARTIAL": "7A5A00",
    "MISSING": "9B1C1C",
    "CONFLICT": "9B1C1C",
    "N_A": "555555",
    "NOT_ANALYZED": "555555",
}


@dataclass(frozen=True, slots=True)
class ReportEvidence:
    document_name: str
    source_locator: str
    excerpt: str
    score: float
    is_manual: bool


@dataclass(frozen=True, slots=True)
class ReportFinding:
    id: str
    requirement_id: str
    candidate_status: str
    status: str
    score: float
    human_verified: bool
    stale: bool
    reviewer: str | None
    owner: str | None
    due_at: datetime | None
    notes: str | None
    evidence_valid: bool
    evidence: tuple[ReportEvidence, ...]


@dataclass(frozen=True, slots=True)
class ReportRequirement:
    id: str
    section: str
    category: str
    mandatory_term: str | None
    obligation_owner: str
    requirement_text: str
    source_document: str
    source_locator: str
    finding: ReportFinding | None


@dataclass(frozen=True, slots=True)
class ReportDocument:
    name: str
    classification: str
    status: str
    volume_name: str | None
    extraction_count: int
    duplicate: bool


@dataclass(frozen=True, slots=True)
class ReportAction:
    id: str
    title: str
    description: str | None
    status: str
    owner: str | None
    due_at: datetime | None
    requirement_id: str | None
    finding_id: str | None


@dataclass(frozen=True, slots=True)
class ComplianceReportData:
    project_id: str
    project_name: str
    solicitation_number: str | None
    agency: str | None
    due_at: datetime | None
    due_timezone: str | None
    sensitivity: str
    generated_at: datetime
    anonymous_public: bool
    analysis_current: bool
    analysis_messages: tuple[str, ...]
    documents: tuple[ReportDocument, ...]
    requirements: tuple[ReportRequirement, ...]
    actions: tuple[ReportAction, ...]


@dataclass(frozen=True, slots=True)
class GapRow:
    rank: int
    priority: str
    requirement: ReportRequirement
    finding: ReportFinding
    linked_actions: tuple[ReportAction, ...]


def _clean_text(value: object | None) -> str:
    if value is None:
        return ""
    return _XML_INVALID.sub("", str(value)).replace("\r\n", "\n").replace("\r", "\n")


def _excerpt(value: object | None, limit: int) -> str:
    text = " ".join(_clean_text(value).split())
    if len(text) <= limit:
        return text
    return text[: max(0, limit - 15)].rstrip() + " ... [truncated]"


def _format_datetime(value: datetime | None) -> str:
    if value is None:
        return "Not set"
    if value.tzinfo is None or value.utcoffset() is None:
        value = value.replace(tzinfo=UTC)
    return value.isoformat(timespec="minutes")


def _display_datetime(value: datetime | None) -> str:
    if value is None:
        return "Not set"
    if value.tzinfo is None or value.utcoffset() is None:
        value = value.replace(tzinfo=UTC)
    return value.astimezone(UTC).strftime("%Y-%m-%d %H:%M UTC")


def _display_date(value: datetime | None) -> str:
    if value is None:
        return "Not set"
    if value.tzinfo is None or value.utcoffset() is None:
        value = value.replace(tzinfo=UTC)
    return value.astimezone(UTC).strftime("%Y-%m-%d")


def _display_project_due(value: datetime | None, timezone_name: str | None) -> str:
    if value is None:
        return "Not set"
    if value.tzinfo is None or value.utcoffset() is None:
        value = value.replace(tzinfo=UTC)
    if timezone_name:
        try:
            local = value.astimezone(ZoneInfo(timezone_name))
        except ZoneInfoNotFoundError:
            pass
        else:
            return f"{local.strftime('%Y-%m-%d %H:%M %Z')} ({timezone_name})"
    return _display_datetime(value)


def _humanize_status(value: str) -> str:
    return value.replace("_", " ").title()


def _sort_datetime(value: datetime | None) -> datetime:
    if value is None:
        return datetime.max.replace(tzinfo=UTC)
    if value.tzinfo is None or value.utcoffset() is None:
        return value.replace(tzinfo=UTC)
    return value.astimezone(UTC)


def _set_run_font(
    run: object,
    *,
    name: str = "Calibri",
    size: float | None = None,
    color: str | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    run.font.name = name  # type: ignore[attr-defined]
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)  # type: ignore[attr-defined]
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)  # type: ignore[attr-defined]
    if size is not None:
        run.font.size = Pt(size)  # type: ignore[attr-defined]
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)  # type: ignore[attr-defined]
    if bold is not None:
        run.bold = bold  # type: ignore[attr-defined]
    if italic is not None:
        run.italic = italic  # type: ignore[attr-defined]


def _remove_children(parent: object, tag: str) -> None:
    for child in list(parent.findall(qn(tag))):  # type: ignore[attr-defined]
        parent.remove(child)  # type: ignore[attr-defined]


def _set_table_borders(table: Table, *, color: str = "D9E2F0") -> None:
    table_properties = table._tbl.tblPr
    _remove_children(table_properties, "w:tblBorders")
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)
        borders.append(border)
    table_properties.append(borders)


def _remove_table_borders(table: Table) -> None:
    table_properties = table._tbl.tblPr
    _remove_children(table_properties, "w:tblBorders")
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "nil")
        borders.append(border)
    table_properties.append(borders)


def _set_cell_margins(cell: _Cell) -> None:
    cell_properties = cell._tc.get_or_add_tcPr()
    _remove_children(cell_properties, "w:tcMar")
    margins = OxmlElement("w:tcMar")
    for edge, value in (("top", 80), ("bottom", 80), ("start", 120), ("end", 120)):
        margin = OxmlElement(f"w:{edge}")
        margin.set(qn("w:w"), str(value))
        margin.set(qn("w:type"), "dxa")
        margins.append(margin)
    cell_properties.append(margins)


def _set_table_geometry(
    table: Table,
    widths_dxa: tuple[int, ...],
    *,
    indent_dxa: int = _TABLE_INDENT_DXA,
) -> None:
    if sum(widths_dxa) != _CONTENT_WIDTH_DXA:
        raise ValueError("Table column widths must total 9360 DXA")
    if any(len(row.cells) != len(widths_dxa) for row in table.rows):
        raise ValueError("Table geometry does not match the number of columns")

    table.autofit = False
    properties = table._tbl.tblPr
    for tag, value in (
        ("w:tblW", _CONTENT_WIDTH_DXA),
        ("w:tblInd", indent_dxa),
    ):
        _remove_children(properties, tag)
        element = OxmlElement(tag)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")
        properties.append(element)
    _remove_children(properties, "w:tblLayout")
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    properties.append(layout)

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        grid_column = OxmlElement("w:gridCol")
        grid_column.set(qn("w:w"), str(width))
        grid.append(grid_column)

    for row_index, row in enumerate(table.rows):
        row_properties = row._tr.get_or_add_trPr()
        no_split = OxmlElement("w:cantSplit")
        row_properties.append(no_split)
        if row_index == 0:
            header = OxmlElement("w:tblHeader")
            header.set(qn("w:val"), "true")
            row_properties.append(header)
        for cell, width in zip(row.cells, widths_dxa, strict=True):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell_properties = cell._tc.get_or_add_tcPr()
            width_element = cell_properties.get_or_add_tcW()
            width_element.set(qn("w:w"), str(width))
            width_element.set(qn("w:type"), "dxa")
            _set_cell_margins(cell)


def _shade_cell(cell: _Cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    _remove_children(properties, "w:shd")
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), fill)
    properties.append(shading)


def _set_cell_text(
    cell: _Cell,
    value: object | None,
    *,
    size: float = 9,
    bold: bool = False,
    color: str = "1F2937",
    align: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.LEFT,
) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.08
    run = paragraph.add_run(_clean_text(value))
    _set_run_font(run, size=size, color=color, bold=bold)


def _style_header_row(table: Table, *, size: float = 8.5) -> None:
    for cell in table.rows[0].cells:
        _shade_cell(cell, "F2F4F7")
        text = cell.text
        _set_cell_text(cell, text, size=size, bold=True, color="0B2545")


def _add_table(
    document: WordDocumentType,
    headers: tuple[str, ...],
    rows: list[tuple[object, ...]],
    widths_dxa: tuple[int, ...],
    *,
    font_size: float = 9,
) -> Table:
    table = document.add_table(rows=1, cols=len(headers))
    for index, header in enumerate(headers):
        _set_cell_text(table.rows[0].cells[index], header, bold=True, color="0B2545")
    for row_data in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row_data):
            align = (
                WD_ALIGN_PARAGRAPH.CENTER
                if isinstance(value, (int, float)) or headers[index] in {"Status", "Section"}
                else WD_ALIGN_PARAGRAPH.LEFT
            )
            _set_cell_text(cells[index], value, size=font_size, align=align)
    _set_table_geometry(table, widths_dxa)
    _set_table_borders(table)
    _style_header_row(table)
    spacer = document.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(4)
    return table


def _configure_styles(document: WordDocumentType) -> None:
    # Use one shared header/footer. Some Word renderers drop content from a
    # duplicated even-page part even when odd/even mode is enabled.
    document.settings.odd_and_even_pages_header_footer = False
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    style_specs = {
        "Normal": (11, "1F2937", 0, 6, 1.1),
        "Heading 1": (16, "2E74B5", 16, 8, 1.0),
        "Heading 2": (13, "2E74B5", 12, 6, 1.0),
        "Heading 3": (12, "1F4D78", 8, 4, 1.0),
    }
    for style_name, (size, color, before, after, line_spacing) in style_specs.items():
        style = document.styles[style_name]
        style.font.name = "Calibri"
        style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
        style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        if style_name != "Normal":
            style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = line_spacing
        if style_name != "Normal":
            style.paragraph_format.keep_with_next = True


def _append_page_field(paragraph: object) -> None:
    label = paragraph.add_run("Page ")  # type: ignore[attr-defined]
    _set_run_font(label, size=8.5, color="667085")
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    field_run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "667085")
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "17")
    properties.extend((color, size))
    text = OxmlElement("w:t")
    text.text = "1"
    field_run.extend((properties, text))
    field.append(field_run)
    paragraph._p.append(field)  # type: ignore[attr-defined]


def _configure_header_footer(document: WordDocumentType, data: ComplianceReportData) -> None:
    section = document.sections[0]
    section.different_first_page_header_footer = False

    header_paragraph = section.header.paragraphs[0]
    header_paragraph.text = ""
    header_paragraph.paragraph_format.space_after = Pt(0)
    run = header_paragraph.add_run("COMPLIANCE ASSESSMENT")
    _set_run_font(run, size=8.5, color="667085", bold=True)
    if data.solicitation_number:
        run = header_paragraph.add_run(f"  |  {_clean_text(data.solicitation_number)}")
        _set_run_font(run, size=8.5, color="667085")

    footer = section.footer
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.text = ""
    footer_paragraph.paragraph_format.space_after = Pt(0)
    footer_table = footer.add_table(rows=1, cols=2, width=Inches(6.5))
    _set_cell_text(
        footer_table.cell(0, 0),
        f"Generated {data.generated_at.astimezone(UTC).strftime('%Y-%m-%d %H:%M UTC')} | "
        f"{data.sensitivity}",
        size=8.5,
        color="667085",
    )
    page_paragraph = footer_table.cell(0, 1).paragraphs[0]
    page_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    page_paragraph.paragraph_format.space_before = Pt(0)
    page_paragraph.paragraph_format.space_after = Pt(0)
    _append_page_field(page_paragraph)
    _set_table_geometry(footer_table, (7000, 2360), indent_dxa=0)
    _remove_table_borders(footer_table)


def _add_labeled_paragraph(
    document: WordDocumentType,
    label: str,
    value: object | None,
    *,
    color: str = "1F2937",
    italic: bool = False,
    after: float = 3,
    keep_with_next: bool = False,
) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.keep_together = True
    paragraph.paragraph_format.keep_with_next = keep_with_next
    label_run = paragraph.add_run(f"{label}: ")
    _set_run_font(label_run, size=10, color="0B2545", bold=True)
    value_run = paragraph.add_run(_clean_text(value) or "Not available")
    _set_run_font(value_run, size=10, color=color, italic=italic)


def _add_callout(
    document: WordDocumentType,
    label: str,
    text: str,
    *,
    fill: str = "F4F6F9",
    accent: str = "0B2545",
) -> None:
    table = document.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    _shade_cell(cell, fill)
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.1
    run = paragraph.add_run(f"{label}: ")
    _set_run_font(run, size=9.5, color=accent, bold=True)
    run = paragraph.add_run(_clean_text(text))
    _set_run_font(run, size=9.5, color="1F2937")
    _set_table_geometry(table, (_CONTENT_WIDTH_DXA,))
    _set_table_borders(table, color="CBD5E1")
    spacer = document.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(4)


def _finding_attention_flags(finding: ReportFinding) -> list[str]:
    flags: list[str] = []
    if finding.stale:
        flags.append("STALE")
    if not finding.evidence_valid:
        flags.append("INVALID EVIDENCE")
    if finding.status != finding.candidate_status and not finding.human_verified:
        flags.append("UNVERIFIED OVERRIDE")
    return flags


def prioritized_gaps(data: ComplianceReportData) -> list[GapRow]:
    actions_by_requirement: defaultdict[str, list[ReportAction]] = defaultdict(list)
    actions_by_finding: defaultdict[str, list[ReportAction]] = defaultdict(list)
    for action in data.actions:
        if action.requirement_id:
            actions_by_requirement[action.requirement_id].append(action)
        if action.finding_id:
            actions_by_finding[action.finding_id].append(action)

    requirements = [
        requirement
        for requirement in data.requirements
        if requirement.finding is not None
        and requirement.finding.status in {"CONFLICT", "MISSING", "PARTIAL"}
    ]
    requirements.sort(
        key=lambda requirement: (
            _STATUS_ORDER[requirement.finding.status],  # type: ignore[union-attr]
            _SECTION_ORDER.get(requirement.section, 4),
            requirement.source_locator.casefold(),
            requirement.id,
        )
    )

    rows: list[GapRow] = []
    for rank, requirement in enumerate(requirements, start=1):
        finding = requirement.finding
        assert finding is not None
        linked = {
            action.id: action
            for action in (
                actions_by_requirement.get(requirement.id, [])
                + actions_by_finding.get(finding.id, [])
            )
        }
        actions = tuple(
            sorted(
                linked.values(),
                key=lambda action: (
                    action.status == "DONE",
                    _sort_datetime(action.due_at),
                    action.id,
                ),
            )
        )
        rows.append(
            GapRow(
                rank=rank,
                priority=_PRIORITY_LABEL[finding.status],
                requirement=requirement,
                finding=finding,
                linked_actions=actions,
            )
        )
    return rows


def _csv_cell(value: object | None) -> str:
    text = _clean_text(value)
    if len(text) > _EXCEL_CELL_LIMIT:
        suffix = " [truncated]"
        text = text[: _EXCEL_CELL_LIMIT - len(suffix)] + suffix
    if text.lstrip().startswith(_FORMULA_PREFIXES):
        text = "'" + text
        if len(text) > _EXCEL_CELL_LIMIT:
            text = text[:_EXCEL_CELL_LIMIT]
    return text


def build_gap_csv(data: ComplianceReportData) -> bytes:
    fieldnames = (
        "rank",
        "priority",
        "status",
        "candidate_status",
        "score",
        "section",
        "category",
        "mandatory_term",
        "requirement_id",
        "requirement",
        "solicitation_source_document",
        "solicitation_source_locator",
        "proposal_evidence_document",
        "proposal_evidence_locator",
        "proposal_evidence_excerpt",
        "evidence_count",
        "evidence_valid",
        "stale",
        "human_verified",
        "reviewer",
        "notes",
        "assigned_owner",
        "requirement_owner",
        "due_at",
        "linked_action_status",
        "linked_action_title",
        "linked_action_owner",
        "linked_action_due_at",
        "project_name",
        "solicitation_number",
        "generated_at",
        "data_boundary",
        "deployment_boundary",
        "analysis_messages",
        "assessment_basis",
    )
    stream = io.StringIO(newline="")
    writer = csv.DictWriter(stream, fieldnames=fieldnames, extrasaction="ignore")
    writer.writeheader()
    generated_at = data.generated_at.astimezone(UTC).isoformat(timespec="seconds")
    for row in prioritized_gaps(data):
        finding = row.finding
        primary_evidence = finding.evidence[0] if finding.evidence else None
        action_text = {
            "status": " | ".join(action.status for action in row.linked_actions),
            "title": " | ".join(action.title for action in row.linked_actions),
            "owner": " | ".join(action.owner or "Unassigned" for action in row.linked_actions),
            "due": " | ".join(_format_datetime(action.due_at) for action in row.linked_actions),
        }
        values: dict[str, object | None] = {
            "rank": row.rank,
            "priority": row.priority,
            "status": finding.status,
            "candidate_status": finding.candidate_status,
            "score": f"{finding.score:.4f}",
            "section": row.requirement.section,
            "category": row.requirement.category,
            "mandatory_term": row.requirement.mandatory_term,
            "requirement_id": row.requirement.id,
            "requirement": row.requirement.requirement_text,
            "solicitation_source_document": row.requirement.source_document,
            "solicitation_source_locator": row.requirement.source_locator,
            "proposal_evidence_document": (
                primary_evidence.document_name if primary_evidence else None
            ),
            "proposal_evidence_locator": (
                primary_evidence.source_locator if primary_evidence else None
            ),
            "proposal_evidence_excerpt": primary_evidence.excerpt if primary_evidence else None,
            "evidence_count": len(finding.evidence),
            "evidence_valid": finding.evidence_valid,
            "stale": finding.stale,
            "human_verified": finding.human_verified,
            "reviewer": finding.reviewer,
            "notes": finding.notes,
            "assigned_owner": finding.owner,
            "requirement_owner": row.requirement.obligation_owner,
            "due_at": _format_datetime(finding.due_at) if finding.due_at else None,
            "linked_action_status": action_text["status"],
            "linked_action_title": action_text["title"],
            "linked_action_owner": action_text["owner"],
            "linked_action_due_at": action_text["due"],
            "project_name": data.project_name,
            "solicitation_number": data.solicitation_number,
            "generated_at": generated_at,
            "data_boundary": data.sensitivity,
            "deployment_boundary": (
                "Synthetic PUBLIC data only; anonymous public deployment."
                if data.anonymous_public
                else f"Recorded project boundary: {data.sensitivity}."
            ),
            "analysis_messages": " | ".join(data.analysis_messages),
            "assessment_basis": "Automated provisional screening; not certified.",
        }
        writer.writerow({field: _csv_cell(values.get(field)) for field in fieldnames})
    return ("\ufeff" + stream.getvalue()).encode("utf-8")


def _effective_status(requirement: ReportRequirement) -> str:
    finding = requirement.finding
    if finding is None:
        return "NOT_ANALYZED"
    if (
        finding.stale
        or not finding.evidence_valid
        or (finding.status != finding.candidate_status and not finding.human_verified)
    ):
        return "NOT_CURRENT"
    return finding.status


def _status_counts(data: ComplianceReportData) -> Counter[str]:
    counts: Counter[str] = Counter()
    for requirement in data.requirements:
        counts[_effective_status(requirement)] += 1
    return counts


def _section_rows(data: ComplianceReportData) -> list[tuple[object, ...]]:
    sections: defaultdict[str, Counter[str]] = defaultdict(Counter)
    for requirement in data.requirements:
        status = _effective_status(requirement)
        sections[requirement.section][status] += 1
        sections[requirement.section]["TOTAL"] += 1
    order = list("ABCDEFGHIJKLM") + ["UNKNOWN"]
    return [
        (
            section,
            sections[section]["TOTAL"],
            sections[section]["COVERED"],
            sections[section]["PARTIAL"],
            sections[section]["MISSING"],
            sections[section]["CONFLICT"],
            sections[section]["N_A"],
            sections[section]["NOT_ANALYZED"] + sections[section]["NOT_CURRENT"],
        )
        for section in order
        if sections[section]["TOTAL"]
    ]


def _add_masthead(document: WordDocumentType, data: ComplianceReportData) -> None:
    title = document.add_paragraph()
    title.paragraph_format.space_before = Pt(8)
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run("PROPOSAL COMPLIANCE ASSESSMENT")
    _set_run_font(run, size=24, color="0B2545", bold=True)

    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(16)
    run = subtitle.add_run(_clean_text(data.project_name))
    _set_run_font(run, size=14, color="475467", bold=True)

    due = _display_project_due(data.due_at, data.due_timezone)
    metadata_rows = [
        ("Solicitation", data.solicitation_number or "Not set"),
        ("Agency", data.agency or "Not set"),
        ("Proposal due", due),
        ("Generated", data.generated_at.astimezone(UTC).strftime("%Y-%m-%d %H:%M UTC")),
        ("Data boundary", data.sensitivity),
        ("Analysis state", "CURRENT" if data.analysis_current else "ATTENTION REQUIRED"),
    ]
    _add_table(
        document,
        ("Project field", "Value"),
        metadata_rows,
        (2300, 7060),
        font_size=9.5,
    )

    boundary = (
        "This export came from an anonymous public deployment. It is visible to any visitor and "
        "must contain only synthetic PUBLIC data. "
        if data.anonymous_public
        else f"Recorded data boundary: {data.sensitivity}. "
    )
    _add_callout(
        document,
        "Automated and provisional",
        boundary
        + "This report is automated screening, not a compliance certification or a substitute "
        "for accountable human proposal review. Scores and statuses may be incomplete or wrong.",
        fill="FFF7E6",
        accent="7A5A00",
    )


def _add_executive_summary(document: WordDocumentType, data: ComplianceReportData) -> None:
    document.add_heading("Executive coverage summary", level=1)
    counts = _status_counts(data)
    total = len(data.requirements)
    compliant = counts["COVERED"] + counts["N_A"]
    coverage = round((compliant / total) * 100, 2) if total else 0.0
    _add_table(
        document,
        ("Requirements", "Coverage", "Covered", "Partial", "Missing", "Conflict", "N/A"),
        [
            (
                total,
                f"{coverage:.2f}%",
                counts["COVERED"],
                counts["PARTIAL"],
                counts["MISSING"],
                counts["CONFLICT"],
                counts["N_A"],
            )
        ],
        (1350, 1350, 1330, 1330, 1330, 1330, 1340),
        font_size=10,
    )
    not_analyzed = counts["NOT_ANALYZED"]
    not_current = counts["NOT_CURRENT"]
    verified = sum(
        requirement.finding is not None and requirement.finding.human_verified
        for requirement in data.requirements
    )
    stale = sum(
        requirement.finding is not None and requirement.finding.stale
        for requirement in data.requirements
    )
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.add_run(
        f"{len(prioritized_gaps(data)):,} findings need proposal attention. "
        f"{not_analyzed:,} requirements have no analysis result; {not_current:,} results are "
        f"excluded from current coverage because they are stale, invalid, or have an unverified "
        f"override; {stale:,} results are stale; "
        f"{verified:,} results carry optional human verification. Verification is not required "
        "one requirement at a time before this report can be used."
    )

    if data.analysis_messages:
        _add_callout(
            document,
            "Input/freshness warning",
            " ".join(data.analysis_messages),
            fill="FDECEC",
            accent="9B1C1C",
        )
    else:
        _add_callout(
            document,
            "Input/freshness check",
            "The report matches the current active solicitation requirements and usable proposal "
            "volumes recorded by the latest analysis run.",
            fill="EDF7F1",
            accent="1F3A5F",
        )

    gap_rows = prioritized_gaps(data)
    blocking = sum(row.finding.status in {"CONFLICT", "MISSING"} for row in gap_rows)
    evaluation_weaknesses = sum(
        row.requirement.section == "M" or row.requirement.category == "EVALUATION_FACTOR"
        for row in gap_rows
    )
    cdrl_exposure = sum(row.requirement.category == "CDRL" for row in gap_rows)
    open_actions = sum(action.status != "DONE" for action in data.actions)
    blocked_actions = sum(action.status == "BLOCKED" for action in data.actions)
    _add_table(
        document,
        ("Review focus", "Current exposure"),
        [
            ("Blocking issues", f"{blocking:,} conflict or missing findings"),
            ("Partial responses", f"{counts['PARTIAL']:,} requirements"),
            ("Evaluation weaknesses", f"{evaluation_weaknesses:,} Section M/evaluation gaps"),
            ("CDRL exposure", f"{cdrl_exposure:,} CDRL-category gaps"),
            ("Contradictions", f"{counts['CONFLICT']:,} conflict findings"),
            ("Action register", f"{open_actions:,} open; {blocked_actions:,} blocked"),
        ],
        (2600, 6760),
        font_size=9.5,
    )
    note = document.add_paragraph()
    note.paragraph_format.space_after = Pt(8)
    run = note.add_run(
        "The application does not yet generate defensible rationale, corrective prose, or cited "
        "recommendations. Those fields remain human-authored rather than fabricated by this report."
    )
    _set_run_font(run, size=9.5, color="667085", italic=True)


def _add_document_scope(document: WordDocumentType, data: ComplianceReportData) -> None:
    document.add_heading("Document scope", level=1)
    rows = [
        (
            item.classification,
            item.name,
            item.status,
            item.volume_name or ("Duplicate" if item.duplicate else ""),
            item.extraction_count,
        )
        for item in data.documents
    ]
    if rows:
        _add_table(
            document,
            ("Role", "Document", "Extraction", "Volume / note", "Characters"),
            rows,
            (1800, 3200, 1300, 1960, 1100),
            font_size=8.5,
        )
    else:
        document.add_paragraph("No documents are recorded for this project.")


def _add_section_summary(document: WordDocumentType, data: ComplianceReportData) -> None:
    document.add_heading("Status by solicitation section", level=1)
    rows = _section_rows(data)
    if not rows:
        document.add_paragraph("No active solicitation requirements are available.")
        return
    _add_table(
        document,
        (
            "Section",
            "Total",
            "Covered",
            "Partial",
            "Missing",
            "Conflict",
            "N/A",
            "Not current / not run",
        ),
        rows,
        (850, 850, 1300, 1200, 1200, 1300, 950, 1710),
        font_size=8.5,
    )


def _detail_requirements(data: ComplianceReportData, gaps: list[GapRow]) -> list[ReportRequirement]:
    if len(data.requirements) <= REPORT_FINDING_LIMIT:
        gap_ids = {row.requirement.id for row in gaps}
        remainder = sorted(
            (item for item in data.requirements if item.id not in gap_ids),
            key=lambda item: (
                _SECTION_ORDER.get(item.section, 4),
                item.source_locator.casefold(),
                item.id,
            ),
        )
        return [row.requirement for row in gaps] + remainder
    return [row.requirement for row in gaps[:REPORT_FINDING_LIMIT]]


def _add_requirement_detail(
    document: WordDocumentType,
    requirement: ReportRequirement,
    index: int,
) -> None:
    finding = requirement.finding
    status = finding.status if finding else "NOT_ANALYZED"
    priority = _PRIORITY_LABEL.get(status, "INFO")
    score = f" | Score {finding.score:.4f}" if finding else ""
    heading = document.add_heading(
        f"{index}. {priority} | {status} | Section {requirement.section}{score}",
        level=3,
    )
    for run in heading.runs:
        run.font.color.rgb = RGBColor.from_string(_STATUS_COLORS.get(status, "1F4D78"))

    _add_labeled_paragraph(
        document,
        "Requirement",
        _excerpt(requirement.requirement_text, 1_200),
        keep_with_next=True,
    )
    _add_labeled_paragraph(
        document,
        "Solicitation source",
        f"{requirement.source_document} | {requirement.source_locator}",
        keep_with_next=True,
    )
    if finding is None:
        _add_labeled_paragraph(
            document,
            "Proposal evidence",
            "No analysis result is available for this requirement.",
            color="9B1C1C",
        )
        return

    if finding.evidence:
        evidence = finding.evidence[0]
        evidence_text = (
            f"{evidence.document_name} | {evidence.source_locator} | "
            f"{_excerpt(evidence.excerpt, 800)}"
        )
        if len(finding.evidence) > 1:
            evidence_text += f" (+{len(finding.evidence) - 1} additional evidence item(s))"
    else:
        evidence_text = "No proposal evidence was attached to this finding."
    _add_labeled_paragraph(
        document,
        "Proposal evidence",
        evidence_text,
        color="9B1C1C" if not finding.evidence_valid else "1F2937",
        keep_with_next=True,
    )

    flags = _finding_attention_flags(finding)
    owner = finding.owner or "Unassigned"
    due = _display_datetime(finding.due_at)
    _add_labeled_paragraph(
        document,
        "Disposition / ownership",
        f"Assigned owner: {owner}; requirement owner: {requirement.obligation_owner}; due: {due}; "
        f"human verified: {'yes' if finding.human_verified else 'no'}"
        + (f"; flags: {', '.join(flags)}" if flags else ""),
        keep_with_next=True,
    )
    _add_labeled_paragraph(
        document,
        "Reviewer notes",
        _excerpt(finding.notes, 700) if finding.notes else "No reviewer notes recorded.",
        italic=not bool(finding.notes),
        after=6,
    )


def _add_findings(document: WordDocumentType, data: ComplianceReportData) -> None:
    document.add_heading("Priority gaps and requirement-level findings", level=1)
    gaps = prioritized_gaps(data)
    details = _detail_requirements(data, gaps)
    if not details:
        document.add_paragraph("No requirement-level findings are available.")
        return

    if len(data.requirements) <= REPORT_FINDING_LIMIT:
        scope_text = f"Showing all {len(details):,} active requirement-level results."
    else:
        scope_text = (
            f"Showing the highest-priority {len(details):,} of {len(gaps):,} proposal gaps across "
            f"{len(data.requirements):,} active requirements. This limit keeps the Word report "
            "reviewable; the complete gap register remains in the companion CSV and the complete "
            "crosswalk remains in the project workbook."
        )
    _add_callout(document, "Detail scope", scope_text)
    for index, requirement in enumerate(details, start=1):
        _add_requirement_detail(document, requirement, index)


def _add_action_register(document: WordDocumentType, data: ComplianceReportData) -> None:
    document.add_heading("Action register", level=1)
    open_actions = sorted(
        (action for action in data.actions if action.status != "DONE"),
        key=lambda action: (
            action.status != "BLOCKED",
            _sort_datetime(action.due_at),
            action.id,
        ),
    )
    if not open_actions:
        document.add_paragraph("No open corrective actions are recorded.")
        return
    shown = open_actions[:REPORT_ACTION_LIMIT]
    rows = [
        (
            _humanize_status(action.status),
            _excerpt(action.title, 180),
            action.owner or "Unassigned",
            _display_date(action.due_at),
            action.requirement_id or action.finding_id or "Project",
        )
        for action in shown
    ]
    _add_table(
        document,
        ("Status", "Corrective action", "Owner", "Due", "Related record"),
        rows,
        (1400, 3300, 1500, 1500, 1660),
        font_size=8.5,
    )
    if len(open_actions) > len(shown):
        document.add_paragraph(
            f"Showing {len(shown):,} of {len(open_actions):,} open actions. Use the live project "
            "action register for the complete current list."
        )


def _add_appendix(document: WordDocumentType, data: ComplianceReportData) -> None:
    document.add_heading("Appendix: report basis and full data", level=1)
    document.add_heading("Status definitions", level=2)
    definitions = [
        ("Covered", "Automated matching found proposal evidence above the covered threshold."),
        ("Partial", "Some matching proposal evidence exists, but coverage is incomplete."),
        ("Missing", "The automated matcher did not find sufficient proposal evidence."),
        ("Conflict", "Matching evidence appears to contradict a required numeric or factual term."),
        ("N/A", "The recorded final disposition marks the requirement not applicable."),
    ]
    _add_table(document, ("Status", "Working definition"), definitions, (1700, 7660), font_size=9)
    document.add_heading("Complete data pointers", level=2)
    _add_labeled_paragraph(
        document,
        "Full gap register",
        "Create reports > Requirements gap report (CSV)",
    )
    _add_labeled_paragraph(
        document,
        "Full project workbook",
        "Raw data exports > Compliance workbook (XLSX)",
    )
    _add_labeled_paragraph(
        document,
        "Live action register",
        "Proposal compliance > Action register",
    )
    _add_callout(
        document,
        "Use limitation",
        "This report records the application's current automated status, evidence, ownership, and "
        "review metadata. It does not assert legal sufficiency, evaluator acceptance, proposal "
        "quality, or final compliance. Re-run analysis after solicitation or proposal changes.",
    )


def build_compliance_report(data: ComplianceReportData) -> bytes:
    document = WordDocument()
    _configure_styles(document)
    _configure_header_footer(document, data)
    properties = document.core_properties
    properties.title = _clean_text(f"Proposal Compliance Assessment - {data.project_name}")
    properties.subject = "Automated provisional solicitation-to-proposal compliance assessment"
    properties.author = "Compliance Review Tool"
    properties.keywords = "proposal, compliance, solicitation, gaps, automated, provisional"
    properties.created = data.generated_at.astimezone(UTC).replace(tzinfo=None)

    _add_masthead(document, data)
    _add_executive_summary(document, data)
    _add_document_scope(document, data)
    _add_section_summary(document, data)
    _add_findings(document, data)
    _add_action_register(document, data)
    _add_appendix(document, data)

    output = io.BytesIO()
    document.save(output)
    return output.getvalue()
