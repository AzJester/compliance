from __future__ import annotations

import csv
import hashlib
import io
from datetime import UTC, datetime

import pytest
from docx import Document as WordDocument
from fastapi import HTTPException
from fastapi.testclient import TestClient

from backend.app.config import Settings
from backend.app.models import (
    CrosswalkFinding,
    CrosswalkRunState,
    CrosswalkStatus,
    Document,
    DocumentStatus,
    ObligationOwner,
    Project,
    ProposalEvidence,
    Requirement,
    RequirementApplicability,
    RequirementCategory,
    RequirementSection,
    Sensitivity,
    ValidationStatus,
)
from backend.app.reports_api import _ensure_export_allowed
from backend.app.workflow_api import _proposal_input_signature, _requirement_input_signature

from .conftest import seed_extracted_document


def _create_project(
    client: TestClient, name: str = "Synthetic Report Project"
) -> dict[str, object]:
    response = client.post(
        "/api/projects",
        json={
            "name": name,
            "solicitation_number": "FAKE-26-R-0042",
            "agency": "Synthetic Agency",
            "due_at": "2026-10-01T17:00:00Z",
            "due_timezone": "UTC",
            "sensitivity": "PUBLIC",
        },
    )
    assert response.status_code == 201, response.text
    return response.json()


def _classify(
    client: TestClient,
    project_id: str,
    document_id: str,
    classification: str,
    *,
    volume_name: str | None = None,
) -> None:
    response = client.patch(
        f"/api/projects/{project_id}/documents/{document_id}/profile",
        json={"classification": classification, "volume_name": volume_name},
    )
    assert response.status_code == 200, response.text


def _word_text(content: bytes) -> str:
    document = WordDocument(io.BytesIO(content))
    paragraphs = [paragraph.text for paragraph in document.paragraphs]
    cells = [cell.text for table in document.tables for row in table.rows for cell in row.cells]
    return "\n".join(paragraphs + cells)


def _seed_analyzed_project(client: TestClient) -> tuple[str, str, str]:
    project = _create_project(client, '../../Café "Report"\r\nInjected')
    project_id = str(project["id"])
    requirement_text = "The offeror shall submit a detailed management plan."
    solicitation_id = seed_extracted_document(
        client,
        project_id,
        f"SECTION L - INSTRUCTIONS\n{requirement_text}",
        name="synthetic-rfp.pdf",
    )
    proposal_id = seed_extracted_document(
        client,
        project_id,
        "Our detailed management plan defines governance, staffing, and controls.",
        name="synthetic-proposal.pdf",
    )
    _classify(client, project_id, solicitation_id, "BASE_SOLICITATION")
    _classify(
        client,
        project_id,
        proposal_id,
        "PROPOSAL_VOLUME",
        volume_name="Volume I - Management",
    )
    extracted = client.post(f"/api/projects/{project_id}/requirements/extract")
    assert extracted.status_code == 200, extracted.text
    requirements = client.get(f"/api/projects/{project_id}/requirements").json()
    assert len(requirements) == 1
    generated = client.post(f"/api/projects/{project_id}/crosswalk/generate")
    assert generated.status_code == 200, generated.text
    finding = client.get(f"/api/projects/{project_id}/crosswalk").json()[0]

    patched_finding = client.patch(
        f"/api/projects/{project_id}/crosswalk/{finding['id']}",
        json={
            "owner": "Management Volume Lead",
            "due_at": "2026-09-15T17:00:00Z",
            "notes": "+Expand the governance evidence.",
        },
    )
    assert patched_finding.status_code == 200, patched_finding.text
    action = client.post(
        f"/api/projects/{project_id}/actions",
        json={
            "title": "Strengthen management-plan evidence",
            "description": "Add a cited governance matrix.",
            "status": "IN_PROGRESS",
            "owner": "Management Volume Lead",
            "due_at": "2026-09-15T17:00:00Z",
            "requirement_id": requirements[0]["id"],
            "finding_id": finding["id"],
        },
    )
    assert action.status_code == 201, action.text
    return project_id, requirements[0]["id"], finding["id"]


def test_word_report_and_gap_csv_are_usable_safe_exports(client: TestClient) -> None:
    project_id, requirement_id, _finding_id = _seed_analyzed_project(client)

    report = client.get(f"/api/projects/{project_id}/exports/compliance-report.docx")
    assert report.status_code == 200, report.text
    assert report.headers["content-type"].startswith(
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    disposition = report.headers["content-disposition"]
    assert "compliance-report.docx" in disposition
    assert "filename*=UTF-8''" in disposition
    assert "\r" not in disposition and "\n" not in disposition
    assert report.headers["cache-control"] == "no-store"
    assert report.headers["x-content-type-options"] == "nosniff"
    word_report = WordDocument(io.BytesIO(report.content))
    assert word_report.settings.odd_and_even_pages_header_footer is False
    section = word_report.sections[0]
    assert "COMPLIANCE ASSESSMENT" in section.header.paragraphs[0].text
    assert section._sectPr.xpath('./w:headerReference[@w:type="even"]') == []
    assert section._sectPr.xpath('./w:footerReference[@w:type="even"]') == []
    text = _word_text(report.content)
    assert "PROPOSAL COMPLIANCE ASSESSMENT" in text
    assert "Automated and provisional" in text
    assert "not a compliance certification" in text
    assert "Synthetic Agency" in text
    assert "Status by solicitation section" in text
    assert "Strengthen management-plan evidence" in text
    assert "Full gap register" in text
    assert "Create reports > Requirements gap report (CSV)" in text

    gaps = client.get(f"/api/projects/{project_id}/exports/gaps.csv")
    assert gaps.status_code == 200, gaps.text
    assert gaps.headers["content-type"].startswith("text/csv")
    assert gaps.content.startswith(b"\xef\xbb\xbf")
    rows = list(csv.DictReader(io.StringIO(gaps.content.decode("utf-8-sig"))))
    assert len(rows) == 1
    row = rows[0]
    assert row["status"] == "PARTIAL"
    assert row["priority"] == "P2"
    assert row["requirement_id"] == requirement_id
    assert row["requirement"] == "The offeror shall submit a detailed management plan."
    assert row["notes"].startswith("'+Expand")
    assert row["solicitation_source_document"] == "synthetic-rfp.pdf"
    assert row["proposal_evidence_document"] == "synthetic-proposal.pdf"
    assert row["assigned_owner"] == "Management Volume Lead"
    assert row["linked_action_status"] == "IN_PROGRESS"
    assert row["linked_action_title"] == "Strengthen management-plan evidence"
    assert row["deployment_boundary"] == "Recorded project boundary: PUBLIC."
    assert row["assessment_basis"] == "Automated provisional screening; not certified."


def test_report_exports_bound_word_detail_but_keep_all_1250_gap_rows(
    client: TestClient,
) -> None:
    project = _create_project(client, "Large Synthetic Compliance Review")
    project_id = str(project["id"])
    solicitation_id = seed_extracted_document(
        client,
        project_id,
        "SECTION M\nSynthetic source",
        name="large-rfp.pdf",
    )
    proposal_id = seed_extracted_document(
        client,
        project_id,
        "Synthetic proposal evidence",
        name="large-proposal.pdf",
    )
    _classify(client, project_id, solicitation_id, "BASE_SOLICITATION")
    _classify(
        client,
        project_id,
        proposal_id,
        "PROPOSAL_VOLUME",
        volume_name="Technical Volume",
    )

    session_factory = client.app.state.session_factory
    with session_factory() as session:
        requirements: list[Requirement] = []
        findings: list[CrosswalkFinding] = []
        evidence: list[ProposalEvidence] = []
        statuses = (
            [CrosswalkStatus.CONFLICT] * 100
            + [CrosswalkStatus.MISSING] * 1_000
            + [CrosswalkStatus.PARTIAL] * 150
        )
        for index, finding_status in enumerate(statuses):
            requirement_id = f"scale-requirement-{index:04d}"
            finding_id = f"scale-finding-{index:04d}"
            section = RequirementSection.M if index % 3 == 0 else RequirementSection.L
            requirements.append(
                Requirement(
                    id=requirement_id,
                    project_id=project_id,
                    document_id=solicitation_id,
                    fingerprint=hashlib.sha256(requirement_id.encode()).hexdigest(),
                    source_text=f"Synthetic requirement source {index}",
                    source_start=index * 10,
                    source_end=(index * 10) + 9,
                    source_locator=f"Page {index // 40 + 1}, item {index + 1}",
                    requirement_text=f"Synthetic requirement {index + 1} shall be addressed.",
                    section=section,
                    category=(
                        RequirementCategory.EVALUATION_FACTOR
                        if section == RequirementSection.M
                        else RequirementCategory.SUBMISSION_INSTRUCTION
                    ),
                    mandatory_term="shall",
                    obligation_owner=ObligationOwner.OFFEROR,
                    applicability=RequirementApplicability.PROPOSAL,
                    confidence=0.95,
                    extraction_method="scale-test",
                    rule_version="scale-test-v1",
                    validation_status=ValidationStatus.PENDING,
                )
            )
            findings.append(
                CrosswalkFinding(
                    id=finding_id,
                    project_id=project_id,
                    requirement_id=requirement_id,
                    candidate_status=finding_status,
                    status=finding_status,
                    score=0.4 if finding_status != CrosswalkStatus.MISSING else 0.0,
                    candidate_signature=hashlib.sha256(finding_id.encode()).hexdigest(),
                    owner="Proposal Writer" if index % 5 == 0 else None,
                    due_at=(datetime(2026, 9, 15, 17, 0, tzinfo=UTC) if index % 7 == 0 else None),
                )
            )
            if finding_status != CrosswalkStatus.MISSING:
                evidence.append(
                    ProposalEvidence(
                        id=f"scale-evidence-{index:04d}",
                        project_id=project_id,
                        finding_id=finding_id,
                        document_id=proposal_id,
                        source_start=0,
                        source_end=18,
                        source_locator="Page 1",
                        excerpt="Synthetic proposal evidence",
                        score=0.4,
                    )
                )
        session.add_all(requirements)
        session.add_all(findings)
        session.add_all(evidence)
        session.flush()
        proposal_document = session.get(Document, proposal_id)
        assert proposal_document is not None
        session.add(
            CrosswalkRunState(
                project_id=project_id,
                requirement_signature=_requirement_input_signature(requirements),
                proposal_signature=_proposal_input_signature([proposal_document]),
                generated_at=datetime(2026, 8, 10, 20, 0, tzinfo=UTC),
            )
        )
        session.commit()

    report = client.get(f"/api/projects/{project_id}/exports/compliance-report.docx")
    assert report.status_code == 200, report.text
    report_text = _word_text(report.content)
    assert "highest-priority 25 of 1,250 proposal gaps" in report_text
    assert "complete gap register remains in the companion CSV" in report_text
    assert len(report.content) < 2_000_000

    gaps = client.get(f"/api/projects/{project_id}/exports/gaps.csv")
    assert gaps.status_code == 200, gaps.text
    rows = list(csv.DictReader(io.StringIO(gaps.content.decode("utf-8-sig"))))
    assert len(rows) == 1_250
    assert [row["status"] for row in rows[:100]] == ["CONFLICT"] * 100
    assert rows[100]["status"] == "MISSING"
    assert rows[-1]["status"] == "PARTIAL"


def test_gap_export_requires_analysis_and_reanalysis_when_findings_are_stale(
    client: TestClient,
) -> None:
    project = _create_project(client, "Current Assessment Guard")
    project_id = str(project["id"])
    solicitation_id = seed_extracted_document(
        client,
        project_id,
        "SECTION L\nThe offeror shall submit a staffing plan.",
        name="guard-rfp.pdf",
    )
    proposal_id = seed_extracted_document(
        client,
        project_id,
        "Our staffing plan identifies roles, qualifications, and coverage.",
        name="guard-proposal.pdf",
    )
    _classify(client, project_id, solicitation_id, "BASE_SOLICITATION")
    _classify(client, project_id, proposal_id, "PROPOSAL_VOLUME", volume_name="Technical")
    assert client.post(f"/api/projects/{project_id}/requirements/extract").status_code == 200

    not_analyzed = client.get(f"/api/projects/{project_id}/exports/gaps.csv")
    assert not_analyzed.status_code == 409
    assert "Analyze the proposal" in not_analyzed.json()["detail"]

    assert client.post(f"/api/projects/{project_id}/crosswalk/generate").status_code == 200
    requirement = client.get(f"/api/projects/{project_id}/requirements").json()[0]
    changed = client.patch(
        f"/api/projects/{project_id}/requirements/{requirement['id']}",
        json={
            "requirement_text": requirement["requirement_text"] + " Include surge coverage.",
            "reviewer": "Synthetic Reviewer",
            "expected_updated_at": requirement["updated_at"],
        },
    )
    assert changed.status_code == 200, changed.text

    stale = client.get(f"/api/projects/{project_id}/exports/gaps.csv")
    assert stale.status_code == 409
    assert "Reanalyze the proposal" in stale.json()["detail"]
    report = client.get(f"/api/projects/{project_id}/exports/compliance-report.docx")
    assert report.status_code == 200
    assert "marked stale and require reanalysis" in _word_text(report.content)


def test_report_displays_due_date_in_the_recorded_timezone(client: TestClient) -> None:
    created = client.post(
        "/api/projects",
        json={
            "name": "Timezone Report",
            "solicitation_number": "FAKE-TZ-001",
            "agency": "Synthetic Agency",
            "due_at": "2026-10-01T23:00:00Z",
            "due_timezone": "America/Phoenix",
            "sensitivity": "PUBLIC",
        },
    )
    assert created.status_code == 201, created.text
    project_id = created.json()["id"]
    report = client.get(f"/api/projects/{project_id}/exports/compliance-report.docx")
    assert report.status_code == 200
    assert "2026-10-01 16:00 MST (America/Phoenix)" in _word_text(report.content)


def test_mixed_usable_and_blank_proposals_keep_the_completed_run_current(
    client: TestClient,
) -> None:
    project = _create_project(client, "Mixed Proposal Scope")
    project_id = str(project["id"])
    solicitation_id = seed_extracted_document(
        client,
        project_id,
        "SECTION L\nThe offeror shall submit a management plan.",
        name="mixed-rfp.pdf",
    )
    valid_id = seed_extracted_document(
        client,
        project_id,
        "Our management plan defines governance, staffing, and delivery controls.",
        name="valid-proposal.pdf",
    )
    blank_id = seed_extracted_document(
        client,
        project_id,
        "Temporary text removed to model a legacy blank extraction.",
        name="blank-proposal.docx",
    )
    _classify(client, project_id, solicitation_id, "BASE_SOLICITATION")
    _classify(client, project_id, valid_id, "PROPOSAL_VOLUME", volume_name="Management")
    _classify(client, project_id, blank_id, "PROPOSAL_VOLUME", volume_name="Appendix")
    with client.app.state.session_factory() as session:
        blank = session.get(Document, blank_id)
        assert blank is not None
        blank.extracted_text = ""
        blank.extraction_count = 0
        blank.status = DocumentStatus.EXTRACTED
        session.commit()

    assert client.post(f"/api/projects/{project_id}/requirements/extract").status_code == 200
    generated = client.post(f"/api/projects/{project_id}/crosswalk/generate")
    assert generated.status_code == 200, generated.text
    assert client.get(f"/api/projects/{project_id}/exports/gaps.csv").status_code == 200
    report = client.get(f"/api/projects/{project_id}/exports/compliance-report.docx")
    assert report.status_code == 200
    text = _word_text(report.content)
    assert "blank, unusable, or duplicate and were not analyzed" in text
    assert "inputs do not match the latest completed analysis run" not in text


def test_anonymous_public_export_guard_rejects_legacy_non_public_project(tmp_path) -> None:
    settings = Settings(
        data_dir=tmp_path / "data",
        frontend_dir=tmp_path,
        deployment_mode="web",
        host="0.0.0.0",
        allowed_origins=("https://compliance.example",),
        trusted_hosts=("compliance.example",),
        managed_proxy=True,
        trust_proxy_headers=True,
        web_access_mode="anonymous",
    )
    project = Project(name="Legacy CUI", sensitivity=Sensitivity.CUI)
    with pytest.raises(HTTPException, match="PUBLIC projects") as exc_info:
        _ensure_export_allowed(project, settings)
    assert exc_info.value.status_code == 403
